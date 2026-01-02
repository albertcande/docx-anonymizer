"""
DOCX Anonymizer - Core Processing Logic

This module provides functionality to anonymize Word documents by replacing
sensitive keywords with placeholder text while preserving document formatting.
"""

import io
import re
import json
import logging
import zipfile
from pathlib import Path
from typing import BinaryIO, Dict, List, Optional, Tuple, Union
from contextlib import contextmanager
from dataclasses import dataclass, field

from docx import Document
from docx.document import Document as DocumentType
from docx.opc.exceptions import PackageNotFoundError
from filelock import FileLock, Timeout

# Configure logging
logger = logging.getLogger(__name__)

# =============================================================================
# Constants
# =============================================================================

DICTIONARY_PATH = Path(__file__).parent / "keyword_dictionary.json"
LOCK_PATH = Path(__file__).parent / "keyword_dictionary.json.lock"

# Input validation limits
MAX_KEYWORD_LENGTH = 200
MAX_KEYWORDS_COUNT = 100
MAX_FILE_SIZE_MB = 50
MAX_FILES_COUNT = 20

# Currency symbols (Unicode escapes for reliability)
CURRENCY_SYMBOLS = (
    r"[\$\u20AC\u00A3\u00A5\u20B9\u20BD\u20BF\u00A2\u20A9\u20AA\u20AB\u0E3F\u20B1\u20B4\u20B8\u20BA\u20BC\u20BE]"
)
FINANCIAL_PATTERN = re.compile(
    rf"({CURRENCY_SYMBOLS}\s?[\d][\d,]*(?:\.\d{{1,2}})?|[\d][\d,]*(?:\.\d{{1,2}})?\s?{CURRENCY_SYMBOLS})",
    re.UNICODE
)

# =============================================================================
# PII Detection Patterns (ordered to avoid overlap - more specific first)
# =============================================================================

PII_PATTERNS = {
    # Credit card first (16 digits) - most specific
    "CREDIT_CARD": re.compile(r'\b(?:\d{4}[-.\s]?){3}\d{4}\b'),
    # Email - very specific pattern
    "EMAIL": re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', re.IGNORECASE),
    # IP address - specific 4-octet pattern
    "IP_ADDRESS": re.compile(r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b'),
    # SSN - 9 digits in specific format (negative lookahead to avoid phone overlap)
    "SSN": re.compile(r'\b(?!\d{3}[-.\s]?\d{3}[-.\s]?\d{4})\d{3}[-.\s]?\d{2}[-.\s]?\d{4}\b'),
    # Phone - last to avoid matching SSN
    "PHONE": re.compile(r'\b(?:\+?1[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}\b'),
    # Date patterns
    "DATE": re.compile(r'\b(?:\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}|\d{4}[/\-]\d{1,2}[/\-]\d{1,2})\b'),
}

# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class ProcessingStats:
    """Statistics from document processing."""
    keywords_replaced: int = 0
    financial_replaced: int = 0
    pii_replaced: Dict[str, int] = field(default_factory=dict)
    
    def total_replacements(self) -> int:
        return self.keywords_replaced + self.financial_replaced + sum(self.pii_replaced.values())

# =============================================================================
# Custom Exceptions
# =============================================================================

class DictionaryLockError(RuntimeError):
    """Raised when dictionary lock cannot be acquired."""
    pass

class DictionarySaveError(IOError):
    """Raised when dictionary save fails."""
    pass

class FileTooLargeError(ValueError):
    """Raised when file exceeds size limit."""
    pass

class TooManyFilesError(ValueError):
    """Raised when too many files are uploaded."""
    pass

# =============================================================================
# Thread-Safe Dictionary Operations
# =============================================================================

@contextmanager
def _dictionary_lock(timeout: float = 5.0):
    """Context manager for thread-safe dictionary access."""
    lock = FileLock(LOCK_PATH, timeout=timeout)
    acquired = False
    try:
        lock.acquire()
        acquired = True
        yield
    except Timeout:
        logger.error("Could not acquire dictionary lock within timeout")
        raise DictionaryLockError("Dictionary is currently locked by another process")
    finally:
        if acquired:
            lock.release()


def _load_dictionary_internal() -> Tuple[Dict[str, str], int]:
    """Internal: Load dictionary from disk. Must be called within lock."""
    if DICTIONARY_PATH.exists():
        try:
            with open(DICTIONARY_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, dict) and "_meta" in data:
                    return data.get("keywords", {}), data.get("_meta", {}).get("next_num", 1)
                elif isinstance(data, dict):
                    return data, len(data) + 1
        except (json.JSONDecodeError, IOError) as e:
            logger.warning(f"Failed to load dictionary: {e}")
    return {}, 1


def _save_dictionary_internal(keywords: Dict[str, str], next_num: int) -> None:
    """Internal: Save dictionary to disk. Must be called within lock. Raises on failure."""
    try:
        data = {"_meta": {"next_num": next_num}, "keywords": keywords}
        with open(DICTIONARY_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    except IOError as e:
        logger.error(f"Failed to save dictionary: {e}")
        raise DictionarySaveError(f"Could not save dictionary: {e}")


def load_dictionary() -> Tuple[Dict[str, str], int]:
    """Load dictionary (thread-safe)."""
    with _dictionary_lock():
        return _load_dictionary_internal()


def get_dictionary_keywords() -> List[str]:
    """Get list of keywords from the persistent dictionary."""
    keywords, _ = load_dictionary()
    return list(keywords.keys())


def validate_keyword(keyword: str) -> str:
    """Validate and sanitize a keyword."""
    if not keyword or not keyword.strip():
        raise ValueError("Keyword cannot be empty")
    keyword = keyword.strip()
    if len(keyword) > MAX_KEYWORD_LENGTH:
        raise ValueError(f"Keyword exceeds {MAX_KEYWORD_LENGTH} characters")
    return keyword


def add_to_dictionary(keywords: List[str], placeholder_template: str = "[REDACTED_{n}]") -> Dict[str, str]:
    """Add new keywords to the persistent dictionary (thread-safe)."""
    with _dictionary_lock():
        dictionary, next_num = _load_dictionary_internal()
        for keyword in keywords:
            try:
                keyword = validate_keyword(keyword)
                if keyword not in dictionary:
                    dictionary[keyword] = placeholder_template.format(n=next_num)
                    next_num += 1
            except ValueError as e:
                logger.warning(f"Skipping invalid keyword: {e}")
        _save_dictionary_internal(dictionary, next_num)
        return dictionary


def clear_dictionary() -> None:
    """Clear all entries from the persistent dictionary (thread-safe)."""
    with _dictionary_lock():
        _save_dictionary_internal({}, 1)


# =============================================================================
# Document Processing
# =============================================================================

def validate_file_size(file_stream: BinaryIO) -> None:
    """Validate file size is within limits."""
    file_stream.seek(0, 2)  # Seek to end
    size_mb = file_stream.tell() / (1024 * 1024)
    file_stream.seek(0)  # Reset to start
    if size_mb > MAX_FILE_SIZE_MB:
        raise FileTooLargeError(f"File exceeds {MAX_FILE_SIZE_MB}MB limit ({size_mb:.1f}MB)")


def _process_table_recursive(table, process_paragraph_fn) -> None:
    """Process a table recursively, including nested tables."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                process_paragraph_fn(para)
            for nested_table in cell.tables:
                _process_table_recursive(nested_table, process_paragraph_fn)


def anonymize_docx(
    file_stream: BinaryIO,
    keywords: Optional[Union[List[str], Dict[str, str]]] = None,
    include_dictionary: bool = True,
    anonymize_financial: bool = False,
    anonymize_pii: bool = False,
    placeholder_template: str = "[REDACTED_{n}]"
) -> Tuple[io.BytesIO, ProcessingStats]:
    """
    Anonymize a Word document by replacing keywords with placeholders.
    
    Raises:
        FileTooLargeError: If file exceeds size limit.
        PackageNotFoundError: If file is not a valid DOCX.
        ValueError: If no anonymization options enabled.
        DictionaryLockError: If dictionary is locked.
    """
    # Validate file size
    validate_file_size(file_stream)
    
    stats = ProcessingStats()
    replacements: Dict[str, str] = {}
    next_placeholder_num = 1
    new_keywords_to_save = []
    
    # Single lock scope for all dictionary operations
    with _dictionary_lock():
        # Load dictionary if enabled
        if include_dictionary:
            dict_entries, next_placeholder_num = _load_dictionary_internal()
            replacements.update(dict_entries)
        
        # Process new keywords
        if keywords:
            if isinstance(keywords, list):
                if len(keywords) > MAX_KEYWORDS_COUNT:
                    raise ValueError(f"Too many keywords (max {MAX_KEYWORDS_COUNT})")
                for kw in keywords:
                    try:
                        kw = validate_keyword(kw)
                        if kw not in replacements:
                            replacements[kw] = placeholder_template.format(n=next_placeholder_num)
                            next_placeholder_num += 1
                            new_keywords_to_save.append(kw)
                    except ValueError as e:
                        logger.warning(f"Skipping invalid keyword: {e}")
            else:
                replacements.update(keywords)
        
        # Save new keywords within same lock scope
        if new_keywords_to_save:
            current_dict, current_num = _load_dictionary_internal()
            for kw in new_keywords_to_save:
                if kw not in current_dict:
                    current_dict[kw] = placeholder_template.format(n=current_num)
                    current_num += 1
            _save_dictionary_internal(current_dict, current_num)
    
    if not replacements and not anonymize_financial and not anonymize_pii:
        raise ValueError("No keywords provided and no anonymization options enabled")
    
    logger.info(f"Processing: {len(replacements)} keywords, financial={anonymize_financial}, pii={anonymize_pii}")
    
    # Load document
    try:
        doc = Document(file_stream)
    except PackageNotFoundError:
        raise ValueError("Invalid or corrupted DOCX file")
    
    # Sort keywords by length (longest first)
    sorted_keywords = sorted(replacements.keys(), key=len, reverse=True)
    patterns = {kw: re.compile(re.escape(kw), re.IGNORECASE) for kw in sorted_keywords}
    
    # Track replacements
    financial_map: Dict[str, str] = {}
    pii_maps: Dict[str, Dict[str, str]] = {k: {} for k in PII_PATTERNS}
    
    def process_text(text: str) -> str:
        if not text:
            return text
        
        # Replace keywords (using subn for accurate count)
        for keyword in sorted_keywords:
            text, count = patterns[keyword].subn(replacements[keyword], text)
            stats.keywords_replaced += count
        
        # Replace PII data
        if anonymize_pii:
            for pii_type, pattern in PII_PATTERNS.items():
                def replace_pii(match, pii_type=pii_type):
                    original = match.group(0)
                    if original not in pii_maps[pii_type]:
                        pii_maps[pii_type][original] = f"[{pii_type}_{len(pii_maps[pii_type]) + 1}]"
                    return pii_maps[pii_type][original]
                text = pattern.sub(replace_pii, text)
        
        # Replace financial data
        if anonymize_financial:
            def replace_financial(match):
                original = match.group(0)
                if original not in financial_map:
                    financial_map[original] = f"[AMOUNT_{len(financial_map) + 1}]"
                return financial_map[original]
            text = FINANCIAL_PATTERN.sub(replace_financial, text)
        
        return text
    
    def process_paragraph(paragraph) -> None:
        for run in paragraph.runs:
            if run.text:
                run.text = process_text(run.text)
    
    # Process document
    for para in doc.paragraphs:
        process_paragraph(para)
    
    for table in doc.tables:
        _process_table_recursive(table, process_paragraph)
    
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                for para in header.paragraphs:
                    process_paragraph(para)
                for table in header.tables:
                    _process_table_recursive(table, process_paragraph)
        
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for para in footer.paragraphs:
                    process_paragraph(para)
                for table in footer.tables:
                    _process_table_recursive(table, process_paragraph)
    
    # Update stats
    stats.financial_replaced = len(financial_map)
    stats.pii_replaced = {k: len(v) for k, v in pii_maps.items() if v}
    
    logger.info(f"Processing complete. Stats: {stats}")
    return _save_to_bytesio(doc), stats


def create_zip_from_files(file_data: List[Tuple[str, io.BytesIO]]) -> io.BytesIO:
    """Create a ZIP file from multiple file data tuples."""
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for filename, data in file_data:
            zip_file.writestr(filename, data.getvalue())
    zip_buffer.seek(0)
    return zip_buffer


def _save_to_bytesio(doc: DocumentType) -> io.BytesIO:
    """Save document to a BytesIO buffer."""
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
